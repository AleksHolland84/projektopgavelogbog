import streamlit as st
from fpdf import FPDF
from datetime import date
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --------------------------------------------------
# Variables
# --------------------------------------------------

# Version number to show ppl on app
version = "0.1"

# --------------------------------------------------
# functions
# --------------------------------------------------
def init_state(key, default):
    if key not in st.session_state:
        st.session_state[key] = default

def generer_word():
    # Funktion til at genere .docx
    doc = Document()

    # Titel
    title = doc.add_heading("Projektopgave – Logbog", level=0)
    #title.alignment = 1  # center

    doc.add_paragraph(f"Dato: {date.today().strftime('%d-%m-%Y')}")

    #doc.add_page_break()

    for dag, indhold in st.session_state.logbog.items():
        if not indhold:
            continue  # spring tomme dage over

        doc.add_heading(dag, level=1)

        for felt, værdi in indhold.items():
            p = doc.add_paragraph()
            p.add_run(f"{felt.capitalize()}:\n").bold = True

            if isinstance(værdi, list):
                if værdi:
                    for item in værdi:
                        doc.add_paragraph(item, style="List Bullet")
                else:
                    doc.add_paragraph("—")
            else:
                doc.add_paragraph(værdi if værdi else "—")

        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --------------------------------------------------
# Grundopsætning
# --------------------------------------------------
st.set_page_config(
    page_title="Projektopgave – Logbog",
    layout="centered"
)

st.title("Projektopgave – Logbog")
st.caption(f"version {version}")


# --------------------------------------------------
# Session state (gemmer elevens svar)
# --------------------------------------------------
if "elev" not in st.session_state:
    st.session_state.elev = {}

if "logbog" not in st.session_state:
    st.session_state.logbog = {
        "Mandag": {},
        "Tirsdag": {},
        "Onsdag": {},
        "Torsdag": {},
        "Fredag": {}
    }


# --------------------------------------------------
# Dilog box
# --------------------------------------------------
intro_key = f"itro_vist"

if intro_key not in st.session_state:
    st.session_state[intro_key] = False


@st.dialog("Lærertips")
def tips():
    st.subheader("Husk:")
    st.markdown("""Logbogen bedømmes ikke på sprog, men på refleksion.
    
Korte, ærlige svar er bedre end “flotte”.""")
    st.markdown("""
Skriv ikke bare hvad du lavede – forklar også hvorfor.
Det er okay at skrive om fejl og problemer.
             """)

    if st.button("Ok"):
        st.session_state[intro_key] = True
        st.rerun()


# Run dialog box
#if not st.session_state[intro_key]:
#    tips()


# --------------------------------------------------
# Trin 1: Elevoplysninger
# --------------------------------------------------
#with st.expander("👤 Elevoplysninger", expanded=True):
 #   st.session_state.elev["navn"] = st.text_input("Navn")
  #  st.session_state.elev["klasse"] = st.text_input("Klasse")
    # st.session_state.elev["gruppe"] = st.text_input("Gruppe")
    #st.session_state.elev["titel"] = st.text_input("Projektets titel")


# --------------------------------------------------
# Trin 2: Vælg dag
# --------------------------------------------------
dag = st.radio(
    "📅 Vælg dag",
    ["Mandag", "Tirsdag", "Onsdag", "Torsdag", "Fredag"],
    horizontal=True
)

st.header(f"Logbog – {dag}")

data = st.session_state.logbog[dag]

# --------------------------------------------------
# Trin 3: Logbogsskema
# --------------------------------------------------
hvor_key = f"hvor_{dag}"
init_state(hvor_key, data.get("hvor", "På skolen"))

valgt_hvor = st.radio(
    "Hvor arbejdede jeg",
    ["På skolen", "Hjemme", "Begge dele"],
    horizontal=True,
    key=hvor_key
)
data["hvor"] = valgt_hvor

st.subheader("A. Hvad har du arbejdet med?")

arbejdsområde_key = f"arbejdsområde_{dag}" 
init_state(arbejdsområde_key, data.get("arbejdsområde", []))

valgte_arbejdsområder = st.multiselect(
    "Vælg én eller flere",
    [
        "Problemformulering",
        "Underspørgsmål",
        "Informationssøgning / kilder",
        "Noter og bearbejdning",
        "Produkt",
        "Samarbejde",
        "Planlægning",
        "Andet"
    ],
    key=arbejdsområde_key
)

data["arbejdsområde"] = valgte_arbejdsområder


st.subheader("B. Konkrete opgaver")
konkrete_key = f"konkrete_{dag}"
init_state(konkrete_key, data.get("konkrete_opgaver", ""))

tekst = st.text_area(
    "Beskriv kort, hvad du konkret har lavet i dag",
    key=konkrete_key,
    height="content"
)

data["konkrete_opgaver"] = tekst



st.subheader("C. Hvordan arbejdede du?")
arbejdsform_key = f"arbejdsform_{dag}"
init_state(arbejdsform_key, data.get("arbejdsform", "Selvstændigt"))
form= st.radio(
    "Vælg det, der passer bedst",
    [
        "Selvstændigt",
        "Sammen med gruppen",
        "Med hjælp fra andre"
    ],
    key=arbejdsform_key
)
data["arbejdsform"] = form



arbejdsmetoder_key = f"arbejdsmetoder_{dag}"
init_state(arbejdsmetoder_key, data.get("arbejdsmetoder", []))

valgte_metoder = st.multiselect(
    "Arbejdsmetoder",
    [
        "Læste og noterede",
        "Sammenlignede kilder",
        "Diskuterede i gruppen",
        "Skrev kladde og rettede",
        "Planlagde næste skridt",
        "Forberedte fremlæggelse"
    ],
    key=arbejdsmetoder_key
)

data["arbejdsmetoder"] = valgte_metoder

st.subheader("D. Refleksion")
fungerede_key = f"fungerede_{dag}"
init_state(fungerede_key, data.get("fungerede", ""))
fungerede = st.text_area(
    "Hvad fungerede godt i dag – og hvorfor?",
    key=fungerede_key,
    height="content"
)
data["fungerede"] = fungerede

svært_key = f"svært_{dag}"
init_state(svært_key, data.get("svært", ""))
svært = st.text_area(
    "Hvad var svært, og hvordan håndterede du det?",
    key=svært_key,
    height="content"
)
data["svært"] = svært

anderledes_key = f"anderledes_{dag}"
init_state(anderledes_key, data.get("anderledes", ""))
anderledes = st.text_area(
    "Hvad ville du gøre anderledes næste gang?",
    key=anderledes_key,
    height="content"
)
data["anderledes"] = anderledes

st.subheader("E. Plan for næste dag")


plan_key = f"plan_{dag}"
init_state(plan_key, data.get("plan", ""))
plan = st.text_area(
    "Hvad er dit næste skridt?",
    key=plan_key,
    height="content"
)
data["plan"] = plan

st.divider()

if st.button("📄 Download logbog som Word"):
    word_file = generer_word()

    st.download_button(
        label="⬇️ Download Word-dokument",
        data=word_file,
        file_name="projektopgave_logbog.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

